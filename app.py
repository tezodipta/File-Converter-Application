from fpdf import FPDF
from io import BytesIO
import pandas as pd
from docx import Document
from flask import Flask, render_template, request, send_file

app = Flask(__name__)

@app.route("/")
def index():
    return render_template("index.html")

# Add a function to configure a font that supports Unicode
def create_unicode_pdf(pdf, text):
    pdf.add_page()
    pdf.add_font('DejaVu', '', './DejaVuSans.ttf', uni=True)  # Load the font
    pdf.set_font('DejaVu', size=12)

    # Add text content
    if isinstance(text, str):
        pdf.multi_cell(0, 10, text)
    elif isinstance(text, list):  # For tabular data
        for line in text:
            pdf.multi_cell(0, 10, line)
    return pdf

def create_dynamic_table_pdf(pdf, data):
    pdf.add_page()
    pdf.add_font('DejaVu', '', './DejaVuSans.ttf', uni=True)  # Load Unicode font
    pdf.set_font('DejaVu', size=8)  # Set font size

    # Table layout
    page_width = pdf.w - 20  # Account for 10-unit margins on both sides
    line_height = 8
    col_widths = []
    num_columns = len(data[0])

    # Calculate column widths proportionally
    for col_index in range(num_columns):
        max_len = max(len(str(row[col_index])) for row in data)
        proportional_width = (max_len / sum(len(str(cell)) for cell in data[0])) * page_width
        col_widths.append(proportional_width)

    # Render the table row by row
    for row in data:
        max_y_in_row = 0  # Tracks maximum height for cells in the row

        for col_index, cell in enumerate(row):
            x_start = pdf.get_x()
            y_start = pdf.get_y()

            # Add the cell content
            pdf.multi_cell(col_widths[col_index], line_height, str(cell), border=1, align='L')

            # Update max_y_in_row
            max_y_in_row = max(max_y_in_row, pdf.get_y())

            # Reset X for the next column
            pdf.set_xy(x_start + col_widths[col_index], y_start)

        # Move to the next row position
        pdf.set_y(max_y_in_row)

        # Page break if needed
        if pdf.get_y() + line_height > pdf.h - 20:  # 20-unit bottom margin
            pdf.add_page()

    return pdf

@app.route("/file_upload", methods=["POST"])
def file_upload():
    input_format = request.form.get("input-format")
    output_format = request.form.get("output-format")
    file = request.files.get("file")

    if not file:
        return "No file uploaded!", 400

    # Check for Excel and CSV formats
    if input_format in ["csv", "excel"]:
        try:
            if input_format == "excel":
                df = pd.read_excel(file)  # Read Excel file
            elif input_format == "csv":
                df = pd.read_csv(file)  # Read CSV file

            df.fillna("", inplace=True)  # Replace NaN values with empty strings

            if output_format == "pdf":
                pdf = FPDF()
                rows = [list(df.columns)] + df.values.tolist()
                rows = [[str(cell) for cell in row] for row in rows]  # Ensure all cells are strings
                pdf = create_dynamic_table_pdf(pdf, rows)  # Use the fixed table creation function

                pdf_output = BytesIO()
                pdf_output.write(pdf.output(dest='S').encode('latin1'))
                pdf_output.seek(0)

                return send_file(pdf_output, download_name="converted.pdf", as_attachment=True)

            elif output_format == "web":
                # Generate an HTML table for web view
                html_table = df.to_html(index=False, escape=False)
                return f"<html><body>{html_table}</body></html>"

        except Exception as e:
            return f"Error processing file: {str(e)}", 400

    # Check for Word format
    elif input_format == "word":
        try:
            doc = Document(file)

            if output_format == "pdf":
                pdf = FPDF()
                paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                pdf = create_unicode_pdf(pdf, paragraphs)

                pdf_output = BytesIO()
                pdf_output.write(pdf.output(dest='S').encode('latin1'))
                pdf_output.seek(0)

                return send_file(pdf_output, download_name="converted.pdf", as_attachment=True)

            elif output_format == "web":
                content = "\n".join(para.text.strip() for para in doc.paragraphs if para.text.strip())
                return f"<pre>{content}</pre>"

        except Exception as e:
            return f"Error processing file: {str(e)}", 400

    return "Unsupported file or conversion type.", 400

if __name__ == "__main__":
    app.run(host='0.0.0.0',debug=True,port=5555)
